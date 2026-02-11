#!/usr/bin/env python3
from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path

from lxml import etree

try:
    from docx import Document
except Exception as e:  # pragma: no cover
    raise SystemExit(f"python-docx is required for this smoke test: {e}")

from refine_docx_from_amended import NS, w_tag, refine_from_amended


def _read_part(docx_path: Path, part: str) -> etree._Element:
    with zipfile.ZipFile(docx_path, "r") as zf:
        return etree.fromstring(zf.read(part))


def main() -> int:
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        original = td_path / "original.docx"
        amended = td_path / "amended.docx"
        refined = td_path / "original_refined.docx"

        doc = Document()
        p = doc.add_paragraph()
        run1 = p.add_run("The emergence of ")
        run1.font.name = "Times New Roman"
        run2 = p.add_run("AI")
        run2.italic = True
        run2.font.name = "Times New Roman"
        run3 = p.add_run(" changes everything.")
        run3.font.name = "Times New Roman"
        doc.save(original)

        doc2 = Document()
        p2 = doc2.add_paragraph()
        p2.add_run("The emergence of ")
        p2.add_run("AI")
        p2.add_run(" fundamentally changes everything.")
        doc2.save(amended)

        changed, skipped = refine_from_amended(original, amended, refined)
        assert changed == 1
        assert skipped == 0

        root = _read_part(refined, "word/document.xml")
        runs = root.xpath("//w:r", namespaces=NS)
        assert runs, "No runs found"

        highlighted_texts = []
        for r in runs:
            rPr = r.find("w:rPr", namespaces=NS)
            if rPr is None:
                continue
            highlight = rPr.find("w:highlight", namespaces=NS)
            if highlight is None:
                continue
            if highlight.get(w_tag("val")) != "yellow":
                continue
            txt = "".join(t.text or "" for t in r.xpath(".//w:t", namespaces=NS))
            highlighted_texts.append(txt)
            assert rPr.find("w:b", namespaces=NS) is not None, "Highlighted run must be bold"

        assert any("fundamentally" in t for t in highlighted_texts), f"Expected changed word highlighted, got: {highlighted_texts}"

        xml = etree.tostring(root, encoding="unicode")
        assert "Times New Roman" in xml, "Expected original font to remain present in document.xml"

    print("OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

